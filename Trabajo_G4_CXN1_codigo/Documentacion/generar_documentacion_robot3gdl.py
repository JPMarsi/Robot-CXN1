from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Documentacion"
DOCX_PATH = OUT_DIR / "Documentacion_Completa_App_Robot_3GDL.docx"
MAP_PATH = OUT_DIR / "Mapa_Proyecto_App_Robot_3GDL.txt"


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph("En Microsoft Word: clic derecho sobre esta tabla y seleccionar Actualizar campo.")


def set_cell_text(cell, text):
    cell.text = ""
    for i, part in enumerate(str(text).split("\n")):
        if i == 0:
            cell.paragraphs[0].add_run(part)
        else:
            cell.add_paragraph(part)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, headers, rows, style="Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return table


def heading(doc, text, level):
    doc.add_heading(text, level=level)


def para(doc, text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Number")


def code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Proyecto Robot antropomorfico 3 GDL - Documentacion tecnica"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Pagina ")
    add_field(footer, "PAGE")
    footer.add_run(" de ")
    add_field(footer, "NUMPAGES")
    return doc


def build_doc():
    doc = setup_document()
    today = datetime.now().strftime("%d/%m/%Y")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documentacion tecnica completa de la aplicacion del robot antropomorfico de 3 GDL")
    run.bold = True
    run.font.size = Pt(20)
    para(doc, "")
    add_table(doc, ["Campo", "Valor"], [
        ["Nombre del proyecto", "Trabajo_G4_CXN1_codigo - Robot antropomorfico 3 GDL"],
        ["Lenguaje y entorno", "MATLAB, interfaz con uifigure/uigridlayout/uitable/uiaxes"],
        ["Fecha de generacion", today],
        ["Version de MATLAB detectada", "24.1.0.2537033 (R2024a), verificada con matlab -batch \"disp(version)\"."],
        ["Archivo principal de inicio", "iniciarRobot3GDL.m"],
        ["Clase principal", "Robot3GDLApp.m, clase Robot3GDLApp < handle"],
        ["Paquete numerico activo", "+robot3gdl"],
    ])
    doc.add_page_break()

    heading(doc, "Indice", 1)
    add_toc(doc)
    doc.add_page_break()

    heading(doc, "1. Objetivo y alcance de la aplicacion", 1)
    para(doc, "La aplicacion fue desarrollada para cargar, validar, calcular, visualizar y animar una rutina de puntos cartesianos para un robot antropomorfico de 3 grados de libertad. La implementacion activa esta compuesta por iniciarRobot3GDL.m, Robot3GDLApp.m y el paquete +robot3gdl. Las carpetas Aparte/Primer modelado y Aparte/Modelos conservan codigo previo y estudios complementarios, pero no son llamadas por la aplicacion nueva.")
    para(doc, "El usuario introduce una tabla de puntos con Nombre, X, Y, Z y Config. Las coordenadas estan en milimetros y Config puede ser auto, arriba o abajo. La aplicacion obtiene configuraciones articulares q1, q2 y q3 en grados, genera una trayectoria articular LSPB por tramos, calcula velocidad y aceleracion articulares, reconstruye posiciones cartesianas por cinematica directa y anima el robot en una vista alambrica y, si existen los archivos configurados, en una vista con modelos STL.")
    add_table(doc, ["Concepto", "Definicion en esta aplicacion"], [
        ["Punto cartesiano", "Fila [X,Y,Z] de TablaRutina, en mm, pasada a robot3gdl.seleccionarSolucion."],
        ["Configuracion articular", "Vector q=[q1,q2,q3] en grados. Se guarda en resultado.Qpuntos y trayectoria.Q."],
        ["Trayectoria cartesiana", "No se interpola directamente en el espacio cartesiano activo. Se reconstruye desde trayectoria.Q con cinematicaDirecta para animacion y error."],
        ["Trayectoria articular", "Matriz Nx3 generada por generarTrayectoriaLSPB a partir de Qpuntos, con densificacion FControl previa."],
        ["Pose del robot", "Conjunto de puntos base-hombro-codo-extremo calculado por cinematicaDirecta; para STL se usan transformaciones jerarquicas."],
        ["Punto alcanzable", "Punto que cumple r no casi nulo, D dentro de [|L2-L3|, L2+L3] y genera soluciones reales."],
        ["Punto valido", "Punto alcanzable cuya solucion elegida queda dentro de qMin/qMax y reconstruye el punto con error <= tolReconstruccion."],
        ["Trayectoria valida", "Trayectoria con tiempos crecientes, muestras finitas, limites articulares, velocidades y aceleraciones dentro de las restricciones codificadas."],
    ])

    heading(doc, "2. Arquitectura general del proyecto", 1)
    para(doc, "La arquitectura activa separa interfaz y calculo. Robot3GDLApp.m administra componentes visuales, callbacks, estado de resultados, graficos y animacion. El paquete +robot3gdl contiene parametros, cinematica, validacion, seleccion de soluciones, generacion LSPB, verificacion de trayectoria e infraestructura STL.")
    add_table(doc, ["Archivo", "Tipo", "Funcion principal", "Entradas", "Salidas", "Es llamado por", "Llama a"], [
        ["iniciarRobot3GDL.m", "Funcion de inicio", "Cierra figuras, limpia clases, agrega la raiz al path e instancia Robot3GDLApp.", "Ninguna", "app", "Usuario MATLAB", "Robot3GDLApp"],
        ["Robot3GDLApp.m", "Clase handle", "Interfaz grafica, callbacks, tablas, graficos, animacion y vista STL.", "Eventos UI y datos de tablas", "Estado app.Resultado y visualizaciones", "iniciarRobot3GDL", "+robot3gdl.*"],
        ["+robot3gdl/parametrosRobot.m", "Funcion", "Define parametros geometricos, articulares, temporales, numericos y STL.", "Ninguna", "struct robot", "Constructor Robot3GDLApp, pruebas", "-"],
        ["+robot3gdl/resolverRutina.m", "Funcion", "Normaliza tabla, resuelve CI punto a punto, valida limites y genera trayectoria.", "tabla, robot, modoSeleccion", "struct resultado", "Robot3GDLApp.validarRutina", "seleccionarSolucion, validarLimites, generarTrayectoriaLSPB"],
        ["+robot3gdl/solucionesCI.m", "Funcion", "Calcula soluciones inversas geometricas y filtra por limites y error.", "P, robot", "soluciones, info", "seleccionarSolucion", "cinematicaDirecta"],
        ["+robot3gdl/seleccionarSolucion.m", "Funcion", "Elige rama manual o automaticamente por costo.", "P, robot, modo, qAnterior, configManual", "q, detalle", "resolverRutina, rutinaEjemplo", "solucionesCI, cinematicaDirecta"],
        ["+robot3gdl/cinematicaDirecta.m", "Funcion", "Calcula base, hombro, codo y extremo.", "q, robot", "matriz 4x3", "CI, animacion, pruebas, resumen", "-"],
        ["+robot3gdl/generarTrayectoriaLSPB.m", "Funcion", "Densifica Q y genera perfil LSPB sincronizado por tramo.", "Qpuntos, robot", "tray", "resolverRutina, pruebas", "resolverRecorridoArticular, validarTrayectoria"],
        ["+robot3gdl/resolverRecorridoArticular.m", "Funcion", "Determina desplazamiento articular admisible sin unwrap global.", "qInicial, qFinal, qMin, qMax, nombre", "struct recorrido", "generarTrayectoriaLSPB", "-"],
        ["+robot3gdl/validarLimites.m", "Funcion", "Verifica matriz Nx3 contra qMin/qMax.", "Q, robot, etiqueta", "reporte o error", "resolverRutina, validarTrayectoria", "-"],
        ["+robot3gdl/validarTrayectoria.m", "Funcion", "Verifica coherencia, limites, velocidades, aceleraciones y paradas.", "tray, robot", "reporte o error", "generarTrayectoriaLSPB", "validarLimites"],
        ["+robot3gdl/transformacionesRobot.m", "Funcion", "Crea transformaciones jerarquicas compatibles con CD para STL.", "q, robot", "struct T", "Robot3GDLApp.actualizarModeloSTL", "makehgtform"],
        ["+robot3gdl/cargarModeloSTL.m", "Funcion", "Carga STL con stlread si existe y valida vertices/caras.", "rutaArchivo, opciones", "modelo", "Robot3GDLApp.inicializarModeloSTL", "stlread"],
        ["tests/ejecutarPruebas.m", "Script/funcion de prueba", "Ejecuta pruebas numericas basicas de rutina, limites, trayectoria y STL.", "Ninguna", "resumen", "Usuario", "+robot3gdl.*"],
        ["Aparte/Primer modelado/*.m", "Codigo previo", "Implementacion anterior de datos, CI/CD, FControl, validaciones, graficos y animacion.", "Variables segun script", "Resultados previos", "No llamado por la app activa", "Funciones locales previas"],
        ["Aparte/Modelos/*.m", "Scripts de estudio", "Volumen de trabajo y singularidades ideales/reales.", "Ninguna o parametros internos", "Figuras", "No llamado por la app activa", "MATLAB grafico"],
    ])
    para(doc, "Flujo real activo: Usuario -> Robot3GDLApp.TablaRutina -> actualizarRobotDesdeUI -> resolverRutina -> normalizarTabla -> seleccionarSolucion -> solucionesCI -> cinematicaDirecta -> validarLimites(Qpuntos) -> generarTrayectoriaLSPB -> densificarPuntosFControl -> resolverRecorridoArticular -> perfilSegmento -> validarTrayectoria -> graficarSeries/prepararEscena3D/precalcularAnimacion -> cinematicaDirecta por muestra -> bucleAnimacion.")
    code(doc, "Usuario -> interfaz -> parametros robot -> tabla rutina -> CI por punto -> seleccion de rama -> validacion articular -> densificacion articular -> LSPB sincronizado -> validacion de trayectoria -> graficos q/dq/ddq -> animacion CD -> vista STL opcional")

    heading(doc, "3. Inicio y funcionamiento general", 1)
    numbered(doc, [
        "El usuario ejecuta app = iniciarRobot3GDL desde MATLAB.",
        "iniciarRobot3GDL cierra figuras, limpia la variable app, limpia clases, ejecuta rehash, agrega la carpeta raiz al path e instancia Robot3GDLApp.",
        "El constructor de Robot3GDLApp carga parametrosRobot, inicializa app.Anim y llama a construirUI.",
        "construirUI crea una uifigure de 1220x760 px y pestanas: Configuracion y rutina, Animacion 3D, Modelo 3D STL, Posicion articular, Velocidad articular y Aceleracion articular.",
        "construirTabConfig crea campos para L1, L2, L3, aceleracion maxima, tiempo nominal, muestreo, tiempo muerto, limites qMin/qMax, modo de seleccion y tablas de rutina/resultados.",
        "construirTabAnim crea ejes 3D y controles Iniciar, Pausar, Continuar, Detener, Reiniciar, Paso render y Retardo.",
        "construirTabModeloSTL crea la vista STL, los checkboxes de ejes/referencia y llama a inicializarModeloSTL.",
        "El constructor llama a cargarEjemplo, que obtiene una rutina validada desde robot3gdl.rutinaEjemplo y la vuelca en TablaRutina.",
    ])
    add_table(doc, ["Accion del usuario", "Componente visual", "Callback o metodo", "Resultado"], [
        ["Cargar ejemplo", "Boton Cargar ejemplo", "app.cargarEjemplo", "Actualiza robot desde UI, carga rutinaEjemplo, limpia resultados y graficos."],
        ["Agregar punto", "Boton Agregar punto", "app.agregarPunto", "Duplica ultimo punto o crea 'Nuevo punto' [675,0,492,'arriba']."],
        ["Eliminar ultimo", "Boton Eliminar ultimo", "app.eliminarPunto", "Elimina la ultima fila si quedan al menos dos puntos; si no, muestra uialert."],
        ["Validar", "Boton Validar", "app.validarRutina", "Ejecuta resolverRutina, guarda app.Resultado y actualiza TablaResultados."],
        ["Calcular", "Boton Calcular", "app.calcular", "Valida, grafica q/dq/ddq, prepara escena 3D y cambia a pestana Posicion articular."],
        ["Iniciar animacion", "Boton Iniciar", "app.iniciarAnimacion", "Valida, grafica, prepara escena, selecciona Animacion 3D y ejecuta bucleAnimacion."],
        ["Pausar", "Boton Pausar", "app.pausarAnimacion", "Establece app.Anim.pausado=true."],
        ["Continuar", "Boton Continuar", "app.continuarAnimacion", "Establece app.Anim.pausado=false."],
        ["Detener", "Boton Detener", "app.detenerAnimacion", "Establece app.Anim.detener=true."],
        ["Reiniciar", "Boton Reiniciar", "app.reiniciarAnimacion", "Detiene, limpia pausa, vuelve a indice/frame 1 y prepara escena."],
        ["Recargar STL", "Boton Recargar STL", "app.inicializarModeloSTL", "Carga modelos STL o usa geometrias provisionales."],
        ["Restablecer camara", "Boton Restablecer camara", "app.restablecerCamaraSTL", "Recalcula limites y vista de la pestana STL."],
    ])

    heading(doc, "4. Validacion de puntos y trayectorias", 1)
    heading(doc, "4.1 Datos de entrada", 2)
    para(doc, "resolverRutina.normalizarTabla acepta table, matriz numerica, celda, string o datos convertibles a celda. Exige al menos dos puntos. Para tablas activas espera Nombre, X, Y, Z y Config. Los campos X/Y/Z se convierten con str2double, admitiendo coma decimal reemplazada por punto; si aparece NaN, Inf o un tipo no convertible se lanza 'La columna X/Y/Z contiene valores no numericos'.")
    para(doc, "Los nombres vacios no invalidan la rutina: se reemplazan por 'Punto i'. Config vacio se convierte en auto. Config solo acepta auto, arriba o abajo; otros valores generan error con la fila.")
    heading(doc, "4.2 Alcanzabilidad geometrica", 2)
    para(doc, "solucionesCI calcula r=hypot(px,py), z=pz-L1, D=hypot(r,z) y q1=atan2d(py,px). Rechaza r < robot.tol porque q1 queda indeterminada. Rechaza D > L2+L3+tol y D < |L2-L3|-tol. No hay una validacion explicita de altura minima/maxima independiente; la altura queda limitada por la distancia D respecto del hombro y por los limites articulares.")
    code(doc, "C = (r^2 + z^2 - L2^2 - L3^2)/(2*L2*L3); C = min(1,max(-1,C)); q3 = +/-acosd(C); q2 = atan2d(z,r) + atan2d(L3*sind(q3), L2 + L3*cosd(q3));")
    heading(doc, "4.3 Restricciones articulares", 2)
    add_table(doc, ["Articulacion", "Limite minimo", "Limite maximo", "Unidad interna", "Unidad mostrada", "Donde se define", "Donde se comprueba"], [
        ["q1 - Base", "-170", "170", "grados", "deg", "parametrosRobot.m robot.qMin/qMax", "solucionesCI, validarLimites, resolverRecorridoArticular"],
        ["q2 - Hombro", "-80", "120", "grados", "deg", "parametrosRobot.m robot.qMin/qMax", "solucionesCI, validarLimites, resolverRecorridoArticular"],
        ["q3 - Codo", "-130", "130", "grados", "deg", "parametrosRobot.m robot.qMin/qMax", "solucionesCI, validarLimites, resolverRecorridoArticular"],
    ])
    para(doc, "solucionesCI marca limitesOk con tolerancia 1e-9 y descarta candidatos que violan limites. resolverRutina vuelve a validar la matriz Qpuntos completa con validarLimites. Despues de generar la trayectoria, validarTrayectoria vuelve a llamar a validarLimites sobre tray.Q.")
    heading(doc, "4.4 Dominio matematico", 2)
    para(doc, "La implementacion activa evita acos fuera de dominio saturando C a [-1,1] despues de las pruebas de alcance. Las divisiones por r usan max(r, robot.tol) al calcular la linea hombro-extremo para clasificar el codo. No se encontraron verificaciones explicitas de resultados complejos en solucionesCI, pero validarTrayectoria si rechaza NaN, Inf o complejos en t, Q, dQ y ddQ.")
    heading(doc, "4.5 Singularidades", 2)
    para(doc, "La aplicacion activa rechaza la singularidad de cintura r < robot.tol. En solucionesCI se calcula un indicador candidato.singularidad = min(abs(sind(q3)), abs(D-(L2+L3))). Este indicador se usa indirectamente en seleccionarSolucion.calcularCostos mediante penalSing = 30/max(abs(sind(q3)),0.05), por lo que las soluciones cerca de q3=0 son penalizadas. No se encontro rechazo explicito por determinante del Jacobiano ni comparacion contra robot.tolSingularidad en el flujo activo. Los scripts Aparte/Modelos/SingularidadesIdeales.m si estudian conceptualmente det(J)=L2*L3*r*sin(theta3), pero no son llamados por la app activa.")
    para(doc, "Esta comprobacion por determinante del Jacobiano es conceptualmente necesaria para analisis dinamico o de velocidad, pero no se encontro implementada explicitamente como criterio de rechazo de trayectoria en el codigo activo.")
    heading(doc, "4.6 Continuidad articular y saltos", 2)
    para(doc, "El codigo activo no aplica unwrap global. resolverRecorridoArticular compara deltaDirecto=qFinal-qInicial con deltaAngularCorto, pero conserva el recorrido directo si el camino corto atravesaria una representacion fuera de los limites mecanicos. Esto evita que un cambio como 160 -> -160 en q1 sea tratado como +40 grados cruzando la zona prohibida fuera de [-170,170]; se usa -320 grados dentro del intervalo y se ajusta el tiempo si hace falta.")
    para(doc, "La seleccion automatica de rama en seleccionarSolucion minimiza un costo compuesto por desplazamiento respecto de qAnterior, penalizacion por cercania a limites, penalizacion por singularidad y penalizacion por cambio de configuracion arriba/abajo. Esto reduce cambios bruscos entre puntos consecutivos, aunque no garantiza continuidad cartesiana entre extremos.")
    heading(doc, "4.7 Validacion de segmentos", 2)
    para(doc, "La aplicacion valida los puntos articulares y la trayectoria articular muestreada. La densificacion se realiza en espacio articular, no cartesiano, con densificarPuntosFControl dentro de generarTrayectoriaLSPB. El criterio usa deltaMaxTeorico = aMax*tfNominal^2/4 y deltaMax = factorDensificacionFControl*deltaMaxTeorico. Si el salto articular maximo de un tramo supera deltaMax, se crean subtramos lineales articulares.")
    para(doc, "No se encontro una verificacion de puntos cartesianos intermedios en segmentos rectos cartesianos, porque el camino activo se diseña en espacio articular. Una trayectoria con extremos cartesianos validos no se valida como segmento cartesiano recto; se valida la trayectoria articular resultante y sus muestras.")
    heading(doc, "4.8 Velocidad y aceleracion", 2)
    para(doc, "A diferencia del primer modelado, la aplicacion activa calcula dQ y ddQ analiticamente dentro de perfilSegmento. validarTrayectoria rechaza aceleraciones mayores que aMax, aceleraciones distintas de -aMax/0/+aMax, velocidades mayores que vMax por articulacion, velocidades inicial/final no nulas y velocidades no nulas en paradas. No se encontro limite de velocidad cartesiana.")
    add_table(doc, ["Verificacion", "Implementada", "Funcion", "Criterio", "Consecuencia si falla"], [
        ["Campos vacios de nombre", "Si", "normalizarTexto", "Nombre vacio -> Punto i", "No falla; se reemplaza."],
        ["Datos no numericos X/Y/Z", "Si", "normalizarNumeros", "~isfinite(valores)", "Error de columna no numerica."],
        ["Cantidad incorrecta de columnas", "Si", "normalizarTabla", "Numerica <3 o celda <5", "Error."],
        ["Minimo de puntos", "Si", "resolverRutina", "n < 2", "Error."],
        ["Puntos repetidos", "No encontrada", "-", "-", "No se rechazan explicitamente."],
        ["Radio r casi cero", "Si", "solucionesCI", "r < robot.tol", "Punto invalido; q1 indeterminada."],
        ["Alcance maximo", "Si", "solucionesCI", "D > L2+L3+tol", "Punto invalido."],
        ["Zona interna no alcanzable", "Si", "solucionesCI", "D < |L2-L3|-tol", "Punto invalido."],
        ["Limites articulares en CI", "Si", "solucionesCI", "qMin/qMax con 1e-9", "Candidato descartado."],
        ["Limites articulares en trayectoria", "Si", "validarLimites", "tray.Q dentro de qMin/qMax", "Error."],
        ["Dominio de acos", "Si", "solucionesCI", "C saturado a [-1,1]", "Evita error numerico tras validar alcance."],
        ["Determinante Jacobiano", "Parcialmente", "seleccionarSolucion/solucionesCI", "Penalizacion por abs(sind(q3)); no det(J)", "Influye costo, no rechaza por si solo."],
        ["Continuidad rama codo", "Si", "seleccionarSolucion", "Costo contra qAnterior y cambio de config", "Elige menor costo."],
        ["Saltos 360/2pi", "Parcialmente", "resolverRecorridoArticular", "No unwrap global; camino corto solo si admisible", "Usa recorrido directo dentro de limites."],
        ["Densificacion de tramos", "Si", "densificarPuntosFControl", "saltoMaximo > factor*aMax*tfNominal^2/4", "Agrega subtramos articulares."],
        ["Velocidad articular maxima", "Si", "validarTrayectoria", "max(abs(dQ(:,j))) <= vMax(j)", "Error."],
        ["Aceleracion articular maxima", "Si", "validarTrayectoria", "abs(ddQ)<=aMax", "Error."],
        ["Velocidad cartesiana maxima", "No encontrada", "-", "-", "No se usa para rechazar."],
    ])

    heading(doc, "5. Flujo completo desde puntos hasta trayectoria", 1)
    numbered(doc, [
        "Lectura de puntos: TablaRutina.Data se convierte a tabla normalizada con Nombre, X, Y, Z y Config.",
        "Conversion de unidades: no hay conversion de mm ni grados en la entrada cartesiana; las coordenadas permanecen en mm y los angulos se calculan en grados con funciones sind/cosd/atan2d/acosd.",
        "Validacion basica: se verifican columnas, cantidad minima, numeros finitos y Config valida.",
        "Cinematica inversa: seleccionarSolucion llama a solucionesCI para cada P_i.",
        "Generacion de soluciones: solucionesCI evalua q3 positivo y negativo, calcula q2, reconstruye con cinematicaDirecta, etiqueta codo y filtra por limites/error.",
        "Seleccion de rama: en modo manual se busca arriba/abajo; en modo auto se minimiza costo con qAnterior.",
        "Validacion de puntos articulares: validarLimites(Q,'Puntos articulares').",
        "Densificacion articular: generarTrayectoriaLSPB inserta puntos si el desplazamiento por subtramo excede el criterio FControl.",
        "Resolucion de recorrido: resolverRecorridoArticular conserva el recorrido fisico dentro de qMin/qMax y documenta si descarta el camino angular corto.",
        "Asignacion de tiempos: cada tramo usa tfNominal o un tf aumentado para cumplir aMax, vMax y tLinealMin.",
        "Interpolacion LSPB: perfilSegmento genera Q, dQ y ddQ con fases de aceleracion, velocidad constante y frenado.",
        "Tiempo muerto: si tMuerto > 0 se agregan muestras de descanso entre tramos, salvo al final si agregarTiempoMuertoFinal=false.",
        "Validacion final: validarTrayectoria revisa consistencia numerica, limites, velocidades, aceleraciones, paradas y eventos de transicion.",
        "Graficos: graficarSeries crea curvas separadas para q, dq y ddq, con marcadores en puntos originales.",
        "Animacion: precalcularAnimacion transforma muestras Q seleccionadas por renderStep en puntos 4x3 con cinematicaDirecta; bucleAnimacion actualiza XData/YData/ZData.",
    ])
    code(doc, "P_i=[x_i,y_i,z_i] mm -> q_i=[q1_i,q2_i,q3_i] deg -> Qpuntos densificados -> Q(t), dQ(t), ddQ(t) -> P_extremo(t)=cinematicaDirecta(Q(t))")

    heading(doc, "6. Modelo geometrico y Denavit-Hartenberg", 1)
    para(doc, "El robot modelado es antropomorfico de 3 GDL con articulaciones rotacionales: q1 gira la cintura alrededor de Z global, q2 gira el hombro en el plano vertical y q3 gira el codo con la convencion de que el angulo absoluto del tercer eslabon es q2-q3. Las longitudes activas por defecto son L1=492 mm, L2=365 mm y L3=310 mm.")
    para(doc, "El codigo no implementa una tabla Denavit-Hartenberg explicita. La cinematica directa esta escrita por ecuaciones geometricas y la visualizacion STL usa makehgtform con transformaciones jerarquicas. Por lo tanto, la tabla siguiente es una reconstruccion equivalente basada en transformacionesRobot.m y cinematicaDirecta.m, no una tabla declarada como tal en el codigo.")
    add_table(doc, ["Articulacion", "theta_i", "d_i", "a_i", "alpha_i", "Variable articular"], [
        ["1", "q1", "0", "0", "0", "q1: giro Z global"],
        ["2", "-q2 en transformaciones STL; q2 en formulas geometricas", "L1", "0", "rotacion equivalente hacia plano vertical", "q2: hombro"],
        ["3", "q3 en transformacion local, equivalente a orientacion q2-q3", "0", "L2", "0", "q3: codo"],
        ["Efector", "0", "0", "L3", "0", "constante"],
    ])
    code(doc, "CD activa:\nx = (L2*cosd(q2) + L3*cosd(q2-q3))*cosd(q1)\ny = (L2*cosd(q2) + L3*cosd(q2-q3))*sind(q1)\nz = L1 + L2*sind(q2) + L3*sind(q2-q3)")
    code(doc, "Transformaciones STL activas:\nT.q1 = Rz(q1)\nT.q2 = Tz(L1)*Ry(-q2)\nT.q3 = Tx(L2)*Ry(q3)\nT.efector = Tx(L3)")

    heading(doc, "7. Cinematica directa", 1)
    para(doc, "robot3gdl.cinematicaDirecta recibe q como vector de 3 elementos en grados y robot con L1, L2, L3. Devuelve una matriz 4x3: P0 base, P1 hombro, P2 codo y P3 extremo. La salida esta en milimetros y se usa para verificar la CI, calcular errores, precalcular animacion, dibujar el robot alambrico y actualizar la referencia STL.")
    add_table(doc, ["Punto", "Formula o calculo", "Uso grafico"], [
        ["P0 Base", "[0,0,0]", "Primer punto del plot3 del robot."],
        ["P1 Hombro", "[0,0,L1]", "Articulacion superior de la columna."],
        ["P2 Codo", "[L2*cosd(q2)*cosd(q1), L2*cosd(q2)*sind(q1), L1+L2*sind(q2)]", "Union entre eslabon 2 y 3."],
        ["P3 Extremo", "P2 + [L3*cosd(q2-q3)*cosd(q1), L3*cosd(q2-q3)*sind(q1), L3*sind(q2-q3)]", "Efector final y traza de trayectoria."],
    ])

    heading(doc, "8. Cinematica inversa", 1)
    para(doc, "La CI activa esta repartida entre solucionesCI y seleccionarSolucion. solucionesCI calcula todas las soluciones reales que cumplen limites y error de reconstruccion. seleccionarSolucion decide cual usar segun modo manual o auto.")
    add_table(doc, ["Solucion", "Caracteristica geometrica", "Ventaja", "Riesgo", "Criterio de seleccion"], [
        ["Codo arriba", "codoValor >= 0 respecto de la recta hombro-extremo en plano r-z", "Puede alejarse de ciertos limites segun punto", "Puede quedar cerca de singularidad o limites", "Manual arriba o menor costo automatico"],
        ["Codo abajo", "codoValor < 0 respecto de la recta hombro-extremo", "Alternativa para puntos altos/bajos", "Cambio de rama puede producir discontinuidad", "Manual abajo o menor costo automatico"],
    ])
    para(doc, "El modo manual se activa con ModoDropDown='manual' y Config de fila distinta de auto. Si la rama solicitada no existe dentro de limites, se lanza error. En modo auto se evalua el costo: desplazamiento, penalLimites, penalSing y cambioConfig. La aplicacion puede cambiar de rama si el costo lo favorece, aunque la penalizacion cambioConfig=80 reduce cambios innecesarios.")

    heading(doc, "9. Diseno de la trayectoria", 1)
    para(doc, "Diseñar la trayectoria en esta app significa transformar una lista de puntos cartesianos validados en una trayectoria articular temporal Q(t), dQ(t), ddQ(t). No se genera un segmento recto cartesiano entre puntos; se conectan configuraciones articulares con perfiles LSPB sincronizados por tramo.")
    add_table(doc, ["Elemento", "Implementacion real"], [
        ["Lista de puntos", "TablaRutina con Nombre, X, Y, Z, Config."],
        ["Camino geometrico", "Camino articular entre Qpuntos; la trayectoria cartesiana se reconstruye por CD."],
        ["Trayectoria con tiempo", "tray.t, tray.Q, tray.dQ, tray.ddQ."],
        ["Muestras de animacion", "Subconjunto de indices 1:renderStep:N, mas ultimo punto."],
        ["Esquinas", "Velocidad articular cero en cada punto objetivo y durante descansos."],
        ["Continuidad de velocidad", "No es continua en aceleracion; dQ se hace cero en paradas. Entre fases LSPB hay cambios de aceleracion."],
    ])
    para(doc, "El metodo de interpolacion real es LSPB/trapezoidal por articulacion. Para cada articulacion activa se usa una fase parabolica de aceleracion +aMax, una fase lineal de velocidad constante y una fase parabolica de frenado -aMax. El tramo se sincroniza con un tiempo comun tf para las tres articulaciones.")
    code(doc, "Fase inicial: q = qi + 0.5*sign(delta)*a*t^2, dq = sign(delta)*a*t, ddq = sign(delta)*a\nFase lineal: q = qi + 0.5*sign(delta)*a*tb^2 + vRegimen*(t-tb), dq=vRegimen, ddq=0\nFase final: q = qf_continuo - 0.5*sign(delta)*a*(tf-t)^2, dq=sign(delta)*a*(tf-t), ddq=-sign(delta)*a")
    para(doc, "El tiempo entre muestras esta definido por tMuestreo=0.02 s. El retardo de reproduccion por defecto en la interfaz es 0.005 s por cuadro renderizado, y renderStep por defecto vale 4, por lo que no necesariamente se dibujan todas las muestras calculadas.")

    heading(doc, "10. Parametros de diseno", 1)
    add_table(doc, ["Parametro", "Simbolo o variable", "Valor actual", "Unidad", "Archivo", "Funcion o propiedad", "Ubicacion aproximada", "Para que sirve", "Efecto de aumentarlo", "Efecto de disminuirlo", "Riesgo"], [
        ["Altura base/hombro", "robot.L1", "492", "mm", "+robot3gdl/parametrosRobot.m", "parametrosRobot", "linea 4", "Desplaza el hombro en Z.", "Eleva el volumen de trabajo.", "Baja el volumen de trabajo.", "Inconsistencia con STL o puntos ejemplo."],
        ["Longitud brazo", "robot.L2", "365", "mm", "+robot3gdl/parametrosRobot.m", "parametrosRobot", "linea 5", "Primer eslabon movil.", "Aumenta alcance y cambia CI.", "Reduce alcance.", "Puntos antes validos pueden fallar."],
        ["Longitud antebrazo", "robot.L3", "310", "mm", "+robot3gdl/parametrosRobot.m", "parametrosRobot", "linea 6", "Segundo eslabon movil.", "Aumenta alcance.", "Amplia zona interna si difiere mas de L2.", "Revisar rutinaEjemplo y STL."],
        ["Limite q1 min/max", "robot.qMin(1)/qMax(1)", "-170/170", "deg", "+robot3gdl/parametrosRobot.m y UI", "parametrosRobot, LimMinFields", "lineas 8-9; UI 102-106", "Restriccion base.", "Permite mas giro.", "Restringe rutinas.", "Camino angular puede cambiar."],
        ["Limite q2 min/max", "robot.qMin(2)/qMax(2)", "-80/120", "deg", "+robot3gdl/parametrosRobot.m y UI", "parametrosRobot, LimMinFields", "lineas 8-9; UI 102-106", "Restriccion hombro.", "Amplia posturas.", "Reduce alcance real.", "CI puede quedarse sin rama."],
        ["Limite q3 min/max", "robot.qMin(3)/qMax(3)", "-130/130", "deg", "+robot3gdl/parametrosRobot.m y UI", "parametrosRobot, LimMinFields", "lineas 8-9; UI 102-106", "Restriccion codo.", "Permite mayor plegado/extension.", "Reduce configuraciones.", "Puede acercar singularidades."],
        ["Aceleracion maxima", "robot.aMax", "100", "deg/s^2", "+robot3gdl/parametrosRobot.m y UI", "FieldAMax", "linea 11; UI 95", "Pendiente de dQ en LSPB.", "Permite mover mas rapido o menos densificar.", "Aumenta tiempos/densificacion.", "Valores altos poco realistas."],
        ["Velocidad maxima", "robot.vMax", "[100 100 100]", "deg/s", "+robot3gdl/parametrosRobot.m", "parametrosRobot", "linea 12", "Tope de velocidad por articulacion.", "Permite perfiles mas rapidos.", "Puede forzar mayor tf o error.", "No esta expuesto en UI."],
        ["Tiempo nominal", "robot.tfNominal", "2.0", "s", "+robot3gdl/parametrosRobot.m y UI", "FieldTfNominal", "linea 13; UI 96", "Duracion deseada por tramo/subtramo.", "Movimiento mas lento; aumenta deltaMax.", "Movimiento mas rapido; mas densificacion.", "Puede violar aMax/vMax si es muy bajo."],
        ["Tramo lineal minimo", "robot.tLinealMin", "0.30", "s", "+robot3gdl/parametrosRobot.m", "parametrosRobot", "linea 14", "Exige meseta visible en LSPB.", "Mas perfil trapezoidal visible.", "Permite perfiles mas triangulares.", "No esta en UI."],
        ["Factor densificacion", "robot.factorDensificacionFControl", "0.90", "-", "+robot3gdl/parametrosRobot.m", "densificarPuntosFControl", "linea 15", "Margen sobre desplazamiento maximo teorico.", "Menos subtramos.", "Mas subtramos.", "Muy alto puede tensionar perfiles."],
        ["Muestreo", "robot.tMuestreo", "0.02", "s", "+robot3gdl/parametrosRobot.m y UI", "FieldTMuestreo", "linea 16; UI 97", "Paso temporal de calculo.", "Menos muestras.", "Mas resolucion.", "Costo de calculo/memoria."],
        ["Tiempo muerto", "robot.tMuerto", "0.2", "s", "+robot3gdl/parametrosRobot.m y UI", "FieldTMuerto", "linea 17; UI 98", "Descanso entre tramos.", "Mas pausa.", "Menos pausa.", "Cambia duracion total."],
        ["Agregar descanso final", "robot.agregarTiempoMuertoFinal", "false", "bool", "+robot3gdl/parametrosRobot.m", "generarTrayectoriaLSPB", "linea 18", "Incluye o no descanso final.", "Agrega pausa final.", "Termina al llegar.", "No expuesto en UI."],
        ["Tolerancia geometrica", "robot.tol", "1e-8", "mm aprox.", "+robot3gdl/parametrosRobot.m", "solucionesCI", "linea 21", "r casi cero y alcance.", "Mas permisivo.", "Mas estricto.", "Puede aceptar singularidad o rechazar puntos validos."],
        ["Tolerancia tiempo", "robot.tolTiempo", "1e-9", "s", "+robot3gdl/parametrosRobot.m", "validarTrayectoria", "linea 22", "Comparacion de tiempos/eventos.", "Fusiona mas tiempos.", "Mas estricto.", "Errores por redondeo."],
        ["Tolerancia velocidad", "robot.tolVelocidad", "1e-7", "deg/s", "+robot3gdl/parametrosRobot.m", "validarTrayectoria", "linea 23", "Velocidad inicial/final cero.", "Mas permisivo.", "Mas estricto.", "Falsos errores numericos."],
        ["Tolerancia reconstruccion", "robot.tolReconstruccion", "1e-5", "mm", "+robot3gdl/parametrosRobot.m", "solucionesCI", "linea 24", "Error CD vs punto deseado.", "Acepta mas error.", "Rechaza por redondeo.", "Perder precision o falsas fallas."],
        ["Paso render", "FieldRenderStep", "4", "muestras/cuadro", "Robot3GDLApp.m", "construirTabAnim/precalcularAnimacion", "lineas 154, 678", "Submuestreo de animacion.", "Menos cuadros; mas rapido.", "Mas cuadros; mas suave.", "Puede ocultar detalle."],
        ["Retardo render", "FieldRenderDelay", "0.005", "s", "Robot3GDLApp.m", "bucleAnimacion", "lineas 156, 634", "Pausa minima entre cuadros.", "Animacion mas lenta.", "Animacion mas rapida.", "Puede saturar UI."],
        ["Limites ejes 3D", "alcance=L2+L3+120", "795", "mm", "Robot3GDLApp.m", "prepararEscena3D/restablecerCamaraSTL", "lineas 570, 588", "Escala visual.", "Mas margen visual.", "Vista mas cerrada.", "Cortar robot o desperdiciar pantalla."],
        ["STL carpeta", "robot.stl.carpeta", "modelos_stl", "ruta", "+robot3gdl/parametrosRobot.m", "parametrosRobot", "linea 30", "Ubicacion modelos STL activos.", "No aplica.", "No aplica.", "Los STL reales estan en STL con nombres Link*.stl; la app busca modelos_stl/eslabon*.stl."],
    ])
    heading(doc, "10.1 Guia rapida de modificaciones solicitadas por el profesor", 2)
    add_table(doc, ["Cambio", "Archivo", "Variable", "Valor actual", "Accion", "Dependencias", "Prueba recomendada"], [
        ["Cambiar longitud de eslabon", "+robot3gdl/parametrosRobot.m o campo UI", "L1/L2/L3", "492/365/310", "Modificar valor y revisar rutina.", "cinematicaDirecta, solucionesCI, transformacionesRobot, STL", "Ejecutar tests/ejecutarPruebas y validar rutina ejemplo."],
        ["Cambiar limite articular", "+robot3gdl/parametrosRobot.m o campo UI", "qMin/qMax", "[-170 -80 -130]/[170 120 130]", "Cambiar limite minimo o maximo.", "validarLimites, solucionesCI, resolverRecorridoArticular", "Probar puntos cerca del limite y trayectoria extendida."],
        ["Aumentar cantidad de puntos de trayectoria", "+robot3gdl/parametrosRobot.m", "tMuestreo", "0.02", "Disminuir tMuestreo.", "validarTrayectoria, animacion", "Verificar tiempo de calculo y suavidad."],
        ["Hacer mas lenta la animacion", "Robot3GDLApp.m o UI", "FieldRenderDelay", "0.005", "Aumentar retardo.", "bucleAnimacion", "Iniciar animacion y observar t label."],
        ["Cambiar tiempo total por tramo", "+robot3gdl/parametrosRobot.m o UI", "tfNominal", "2.0", "Aumentar/disminuir.", "densificacion, LSPB, tiemposMinimos", "Comparar tray.t(end), max dQ y max ddQ."],
        ["Aumentar precision geometrica", "+robot3gdl/parametrosRobot.m", "tolReconstruccion", "1e-5", "Disminuir con cuidado.", "solucionesCI", "Validar rutina ejemplo."],
        ["Cambiar limites graficos", "Robot3GDLApp.m", "alcance y zlim", "L2+L3+120", "Editar margen 120.", "Animacion 3D y STL", "Preparar escena y revisar encuadre."],
        ["Cambiar unidades visualizadas", "Robot3GDLApp.m y funciones", "Etiquetas mm/deg", "mm, deg", "No basta cambiar textos: convertir calculos.", "Todas las funciones cinematicas", "Prueba de reconstruccion CD/CI."],
    ])

    heading(doc, "11. Diccionario de variables e interfaz", 1)
    add_table(doc, ["Texto mostrado", "Variable interna", "Definicion", "Como se calcula", "Unidad", "Momento de actualizacion", "Interpretacion"], [
        ["L1", "app.FieldL1 / robot.L1", "Altura al hombro.", "parametrosRobot o valor UI.", "mm", "Constructor y actualizarRobotDesdeUI.", "Desplazamiento vertical base-hombro."],
        ["L2", "app.FieldL2 / robot.L2", "Longitud brazo.", "parametrosRobot o valor UI.", "mm", "Constructor y actualizarRobotDesdeUI.", "Distancia hombro-codo."],
        ["L3", "app.FieldL3 / robot.L3", "Longitud antebrazo.", "parametrosRobot o valor UI.", "mm", "Constructor y actualizarRobotDesdeUI.", "Distancia codo-efector."],
        ["Aceleracion max.", "app.FieldAMax / robot.aMax", "Aceleracion angular maxima.", "Valor UI.", "deg/s^2", "Antes de validar/calcular.", "Pendiente de velocidad LSPB."],
        ["Tiempo nominal", "app.FieldTfNominal / robot.tfNominal", "Tiempo deseado por tramo.", "Valor UI.", "s", "Antes de validar/calcular.", "Duracion base del perfil."],
        ["Muestreo", "app.FieldTMuestreo / robot.tMuestreo", "Paso temporal.", "Valor UI.", "s", "Antes de validar/calcular.", "Resolucion de Q/dQ/ddQ."],
        ["Tiempo muerto", "app.FieldTMuerto / robot.tMuerto", "Pausa entre tramos.", "Valor UI.", "s", "Antes de validar/calcular.", "Descanso con velocidad cero."],
        ["Limites q min/max", "app.LimMinFields/app.LimMaxFields", "Restricciones articulares.", "Valores UI.", "deg", "Antes de validar/calcular.", "Rango mecanico permitido."],
        ["Nombre, X, Y, Z, Config", "app.TablaRutina.Data", "Rutina cartesiana.", "Usuario o rutinaEjemplo.", "mm para X/Y/Z", "Carga/edicion/validacion.", "Puntos deseados y rama."],
        ["q1/q2/q3 [deg]", "app.TablaResultados.Data", "Solucion articular por punto.", "resultado.Qpuntos.", "deg", "actualizarTablaResultados.", "Configuracion seleccionada."],
        ["Config", "resultado.detalles(i).config", "Rama de codo detectada/elegida.", "codoValor en CI.", "-", "Validacion.", "arriba o abajo."],
        ["Error [mm]", "resultado.detalles(i).error", "Norma de Pcalc-P.", "cinematicaDirecta(q)-P.", "mm", "Validacion.", "Error de reconstruccion."],
        ["Paso render", "app.FieldRenderStep", "Submuestreo para animacion.", "round(valor UI).", "muestras", "precalcularAnimacion.", "Cuantas muestras se saltan por cuadro."],
        ["Retardo [s]", "app.FieldRenderDelay", "Espera minima entre cuadros.", "Valor UI.", "s", "bucleAnimacion.", "Velocidad visual."],
        ["t = ...", "app.TiempoLabel", "Tiempo y frame actual.", "tray.t(i), frame, nFrames.", "s", "actualizarFrame.", "Progreso de animacion."],
        ["Posicion articular", "app.AxQ", "Graficos q(t).", "tray.Q.", "deg", "graficarSeries.", "Posicion por articulacion."],
        ["Velocidad articular", "app.AxDQ", "Graficos dq(t).", "tray.dQ.", "deg/s", "graficarSeries.", "Velocidad angular."],
        ["Aceleracion articular", "app.AxDDQ", "Graficos ddq(t).", "tray.ddQ.", "deg/s^2", "graficarSeries.", "Aceleracion angular."],
    ])

    heading(doc, "12. Representacion grafica, animacion y STL", 1)
    para(doc, "La animacion alambrica usa puntos calculados por cinematicaDirecta. prepararEscena3D dibuja los puntos de la rutina como marcadores negros, el robot como linea con marcadores y la traza del efector. precalcularAnimacion calcula puntosRobot de dimension 4x3xnFrames y extremos de dimension nFramesx3. actualizarFrame actualiza XData, YData y ZData sin recrear la figura.")
    para(doc, "La vista STL crea una jerarquia hgtransform: base, q1, eslabon1, q2, eslabon2, q3, eslabon3 y efector. La app busca archivos en modelos_stl con nombres base.stl, eslabon1.stl, eslabon2.stl, eslabon3.stl y efector.stl. En el proyecto existen STL/base.stl y STL/Link1.stl a Link3.stl, pero no coinciden con la carpeta y nombres activos. Si no se encuentran, la app usa prismas provisionales con crearPrismaEslabon.")
    para(doc, "La carpeta modelos_stl incluye LEEME_STL.txt con instrucciones para exportar desde Fusion, usar milimetros, copiar archivos, revisar nombres en parametrosRobot, recargar la pestana Modelo 3D STL y ajustar escala/rotacion/traslacion local.")

    heading(doc, "13. Pruebas y verificacion", 1)
    para(doc, "tests/ejecutarPruebas.m contiene pruebas de rutina ejemplo, ajuste automatico de tiempo LSPB, rechazo de punto invalido, recorridos articulares y limites, trayectoria articular extendida e infraestructura STL sin archivos. Las pruebas usan asserts sobre error de reconstruccion, NaN/Inf, violaciones de limites, velocidad inicial/final cero, vMax, aMax, densificacion y consistencia de transformaciones STL con cinematica directa.")
    para(doc, "La suite fue ejecutada desde este entorno con MATLAB R2024a mediante matlab -batch. Resultado: 6 pruebas superadas, 0 fallidas. Metricas reportadas: error maximo de reconstruccion 1.27106e-13 mm, sin NaN/Inf, sin violaciones de limites, velocidad maxima [63.4813, 22.5869, 32.9254] deg/s y aceleracion maxima [100, 100, 100] deg/s^2.")
    code(doc, "addpath(fullfile(pwd, 'tests'))\nresumen = ejecutarPruebas;")

    heading(doc, "14. Limitaciones del analisis", 1)
    bullets(doc, [
        "No se abrio ni descompilo Robot3GDL_App.mlapp; se documento como implementacion previa conservada en Aparte/Primer modelado. El flujo activo usa Robot3GDLApp.m.",
        "No se analizaron internamente los PDF y DOCX de Aparte/Documentos porque son material teorico complementario y no intervienen en la ejecucion de la app activa.",
        "Los archivos .fig y .jpg de Aparte/Imagenes se identificaron como recursos de estudio, pero no son llamados por la aplicacion activa.",
        "Los STL reales presentes en STL no coinciden con la ruta/nombres configurados en parametrosRobot.m; la documentacion refleja esta diferencia sin modificar archivos.",
        "No se encontro implementacion activa de validacion por determinante del Jacobiano ni limite de velocidad cartesiana.",
    ])

    heading(doc, "15. Consecuencias de modificar funciones principales", 1)
    add_table(doc, ["Funcion", "Consecuencia de modificarla"], [
        ["parametrosRobot", "Cambia comportamiento global de alcance, limites, trayectoria, tolerancias y STL. Requiere correr pruebas."],
        ["cinematicaDirecta", "Afecta verificacion de CI, animacion, errores, transformaciones y pruebas; debe mantenerse compatible con solucionesCI."],
        ["solucionesCI", "Afecta aceptacion/rechazo de puntos y ramas disponibles. Un cambio puede validar puntos imposibles o rechazar puntos validos."],
        ["seleccionarSolucion", "Cambia continuidad entre puntos, preferencia arriba/abajo y cercania a limites/singularidades."],
        ["resolverRecorridoArticular", "Cambia la interpretacion de saltos angulares y el respeto de limites mecanicos durante tramos."],
        ["generarTrayectoriaLSPB", "Afecta tiempos, velocidades, aceleraciones, densificacion y paradas. Es la funcion mas sensible del control temporal."],
        ["validarTrayectoria", "Modificarla cambia que trayectorias son aceptadas como fisicamente coherentes."],
        ["Robot3GDLApp", "Cambios pueden afectar interfaz, callbacks, estado app.Resultado, graficos y animacion."],
    ])

    doc.save(DOCX_PATH)


def build_map():
    files = [
        ("iniciarRobot3GDL.m", "Entrada unica. Instancia Robot3GDLApp."),
        ("Robot3GDLApp.m", "Clase UI, callbacks, tablas, graficos, animacion y STL."),
        ("+robot3gdl/parametrosRobot.m", "Parametros del robot, limites, tiempos, tolerancias y STL."),
        ("+robot3gdl/resolverRutina.m", "Normalizacion de tabla, CI punto a punto, validacion y trayectoria."),
        ("+robot3gdl/seleccionarSolucion.m", "Seleccion manual/automatica de rama de CI."),
        ("+robot3gdl/solucionesCI.m", "Cinematica inversa y filtros de alcance, limites y error."),
        ("+robot3gdl/cinematicaDirecta.m", "CD base-hombro-codo-extremo."),
        ("+robot3gdl/generarTrayectoriaLSPB.m", "Densificacion articular, LSPB, tiempos y validacion."),
        ("+robot3gdl/resolverRecorridoArticular.m", "Recorrido angular admisible dentro de limites."),
        ("+robot3gdl/validarLimites.m", "Chequeo qMin/qMax."),
        ("+robot3gdl/validarTrayectoria.m", "Chequeo integral de t, Q, dQ, ddQ."),
        ("+robot3gdl/transformacionesRobot.m", "Transformaciones jerarquicas para STL."),
        ("+robot3gdl/cargarModeloSTL.m", "Carga STL y valida triangulacion."),
        ("+robot3gdl/construirTransformacionSTL.m", "Transformacion local escala/rotacion/traslacion."),
        ("+robot3gdl/crearPrismaEslabon.m", "Geometria provisional rectangular."),
        ("+robot3gdl/rutinaEjemplo.m", "Rutina de demostracion validada."),
        ("+robot3gdl/resumenPruebas.m", "Metricas numericas sobre resultado."),
        ("tests/ejecutarPruebas.m", "Suite de pruebas numericas basicas."),
        ("Aparte/Primer modelado", "Implementacion anterior no llamada por la app activa."),
        ("Aparte/Modelos", "Scripts de volumen de trabajo y singularidades."),
        ("modelos_stl/LEEME_STL.txt", "Instrucciones para modelos STL activos."),
        ("STL/*.stl", "Modelos STL existentes con nombres no configurados por defecto."),
    ]
    lines = [
        "MAPA DEL PROYECTO - APP ROBOT 3 GDL",
        f"Raiz: {ROOT}",
        f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",
        "",
        "Flujo activo:",
        "iniciarRobot3GDL -> Robot3GDLApp -> resolverRutina -> seleccionarSolucion -> solucionesCI -> cinematicaDirecta -> validarLimites -> generarTrayectoriaLSPB -> validarTrayectoria -> graficos/animacion",
        "",
        "Archivos y responsabilidades:",
    ]
    for path, desc in files:
        lines.append(f"- {path}: {desc}")
    lines.extend([
        "",
        "Parametros principales:",
        "L1=492 mm, L2=365 mm, L3=310 mm",
        "qMin=[-170 -80 -130] deg, qMax=[170 120 130] deg",
        "aMax=100 deg/s^2, vMax=[100 100 100] deg/s",
        "tfNominal=2.0 s, tMuestreo=0.02 s, tMuerto=0.2 s",
        "tol=1e-8, tolTiempo=1e-9, tolVelocidad=1e-7, tolReconstruccion=1e-5",
        "",
        "Limitaciones registradas:",
        "- MATLAB detectado: 24.1.0.2537033 (R2024a). Suite tests/ejecutarPruebas: 6 superadas, 0 fallidas.",
        "- Robot3GDL_App.mlapp no fue inspeccionado internamente.",
        "- PDF/DOCX/FIG/JPG complementarios no intervienen en el flujo activo.",
    ])
    MAP_PATH.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    OUT_DIR.mkdir(exist_ok=True)
    build_doc()
    build_map()
    print(DOCX_PATH)
    print(MAP_PATH)
